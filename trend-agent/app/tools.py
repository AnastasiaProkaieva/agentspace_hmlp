import base64
import json
from pathlib import Path
import requests
from email.message import EmailMessage
from fpdf import FPDF
from docx import Document

from google.cloud import storage
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload
from google.adk.tools import ToolContext
from typing import Dict, Any

from .config import settings
from .auth_tools import get_authenticated_credentials

REPORT_TEXT_STATE_KEY = "current_report_draft"
VERIFICATION_REASONS_KEY = "verification_reasons"


def classify_topic(topic: str) -> Dict[str, str]:
    sensitive_keywords = [
        "politics",
        "conflict",
        "finance",
        "medical",
        "danger",
        "war",
        "religion",
    ]
    if any(keyword in topic.lower() for keyword in sensitive_keywords):
        return {
            "classification": "sensitive",
            "reason": "The topic is sensitive and I cannot generate a market report on it.",
        }
    return {"classification": "safe"}


def check_report_for_verification(
    report_text: str, tool_context: ToolContext
) -> Dict[str, Any]:
    cleaned_report_text = report_text.replace('"', "")
    # Save the cleaned text to the state for the next tool to use.
    tool_context.state[REPORT_TEXT_STATE_KEY] = cleaned_report_text

    reasons = []
    mentioned_competitors = []
    if len(cleaned_report_text) > settings.REPORT_CHAR_LIMIT_FOR_REVIEW:
        reasons.append(
            f"Report length ({len(cleaned_report_text)} chars) exceeds the limit of {settings.REPORT_CHAR_LIMIT_FOR_REVIEW}."
        )
    for competitor in settings.COMPETITORS:
        if competitor in cleaned_report_text.lower():
            mentioned_competitors.append(competitor)
    if mentioned_competitors:
        reasons.append(
            f"Report mentions competitor(s): {', '.join(list(set(mentioned_competitors)))}."
        )
    if reasons:
        return {"verification_needed": True, "reasons": reasons}
    else:
        return {
            "verification_needed": False,
            "reasons": ["Report is within guidelines."],
        }


def submit_report_for_verification(tool_context: ToolContext) -> Dict[str, Any]:

    report_to_review = tool_context.state.get(REPORT_TEXT_STATE_KEY)
    if not report_to_review:
        return {
            "error": "No report found in state. 'check_report_for_verification' must be called first.",
            "status": "failed",
        }
    # Call the helper with initiate_auth_flow=False.
    # We expect auth to already be done. If not, this tool should fail clearly.
    creds = get_authenticated_credentials(tool_context, initiate_auth_flow=False)
    if not creds:
        return {
            "status": "failed_auth",
            "message": "Authentication is missing. The `authenticate_google_services` tool MUST be called successfully before using this tool.",
        }
    try:
        response = requests.post(
            f"{settings.REVIEW_APP_BASE_URL}/reviews",
            json={"outline": report_to_review},
        )
        response.raise_for_status()
        data = response.json()
        review_id = data["review_id"]
        review_url = f"{settings.REVIEW_APP_BASE_URL}/reviews/{review_id}/view"
        gmail_service = build("gmail", "v1", credentials=creds)
        message = EmailMessage()
        message.set_content(
            f"""A new market analysis report requires your verification.\n\nPlease review it here: {review_url}"""
        )
        message["To"] = settings.REVIEWER_EMAIL
        message["Subject"] = f"""Market Report Needs Verification (ID: {review_id})"""
        encoded_message = base64.urlsafe_b64encode(message.as_bytes()).decode()
        gmail_service.users().messages().send(
            userId="me", body={"raw": encoded_message}
        ).execute()
        return {
            "review_id": review_id,
            "review_url": review_url,
            "status": "pending_report_approval",
        }
    except Exception as e:
        return {"error": str(e), "status": "failed"}


def save_report_formats(
    report_markdown: str, report_title: str, tool_context: ToolContext
) -> Dict[str, Any]:
    creds = get_authenticated_credentials(tool_context, initiate_auth_flow=False)
    if not creds:
        return {
            "status": "failed_auth",
            "message": "Authentication is missing. The `authenticate_google_services` tool MUST be called successfully before using this tool.",
        }
    try:
        script_dir = Path(__file__).parent.resolve()
        font_path = script_dir / "OpenSans.ttf"
        work_dir = Path("./temp_reports")
        work_dir.mkdir(exist_ok=True)
        title_slug = report_title.lower().replace(" ", "_").replace("'", "")
        md_path = work_dir / f"{title_slug}.md"
        pdf_path = work_dir / f"{title_slug}.pdf"
        docx_path = work_dir / f"{title_slug}.docx"
        md_path.write_text(report_markdown, encoding="utf-8")
        pdf = FPDF()
        pdf.add_font("CustomFont", "", str(font_path), uni=True)
        pdf.set_font("CustomFont", "", 12)
        pdf.add_page()
        pdf.multi_cell(0, 10, report_markdown)
        pdf.output(pdf_path)
        doc = Document()
        doc.add_heading(report_title, 0)
        doc.add_paragraph(report_markdown)
        doc.save(docx_path)
        clean_bucket_name = settings.STAGING_BUCKET.replace("gs://", "")
        gcs_client = storage.Client(credentials=creds)
        bucket = gcs_client.bucket(clean_bucket_name)
        drive_service = build("drive", "v3", credentials=creds)
        gcs_links, drive_links = {}, {}

        for file_path in [md_path, pdf_path, docx_path]:
            blob = bucket.blob(f"reports/{file_path.name}")
            blob.upload_from_filename(str(file_path))
            gcs_links[file_path.suffix] = blob.public_url
            file_metadata = {"name": file_path.name}
            media = MediaFileUpload(
                str(file_path), mimetype="application/octet-stream", resumable=True
            )
            file = (
                drive_service.files()
                .create(body=file_metadata, media_body=media, fields="webViewLink")
                .execute()
            )
            drive_links[file_path.suffix] = file.get("webViewLink")
        return {"status": "success", "gcs_urls": gcs_links, "drive_urls": drive_links}
    except Exception as e:
        return {
            "status": "error",
            "message": f"Failed to save report formats: {str(e)}",
        }
